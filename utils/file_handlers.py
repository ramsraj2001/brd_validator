"""File handling utilities for BRD documents."""

import streamlit as st
import tempfile
import os
from pathlib import Path
from typing import Optional, Dict, Any
import docx
import PyPDF2
import json

from config.settings import SUPPORTED_FORMATS, MAX_FILE_SIZE

class FileHandler:
    """Handles file upload and processing operations."""
    
    def __init__(self):
        self.temp_files = []
    
    def validate_file(self, uploaded_file) -> Dict[str, Any]:
        """Validate uploaded file format and size."""
        validation_result = {
            'valid': False,
            'error': None,
            'file_info': {}
        }
        
        if not uploaded_file:
            validation_result['error'] = "No file uploaded"
            return validation_result
        
        # Check file size
        if uploaded_file.size > MAX_FILE_SIZE:
            validation_result['error'] = f"File size ({uploaded_file.size / 1024 / 1024:.1f}MB) exceeds limit ({MAX_FILE_SIZE / 1024 / 1024:.1f}MB)"
            return validation_result
        
        # Check file format
        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in SUPPORTED_FORMATS:
            validation_result['error'] = f"Unsupported file format: {file_extension}. Supported formats: {', '.join(SUPPORTED_FORMATS)}"
            return validation_result
        
        validation_result['valid'] = True
        validation_result['file_info'] = {
            'name': uploaded_file.name,
            'size': uploaded_file.size,
            'type': uploaded_file.type,
            'extension': file_extension
        }
        
        return validation_result
    
    def save_temp_file(self, uploaded_file) -> Optional[str]:
        """Save uploaded file to temporary location."""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                temp_path = tmp_file.name
                self.temp_files.append(temp_path)
                return temp_path
        except Exception as e:
            st.error(f"Error saving file: {str(e)}")
            return None
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text content from PDF file."""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text_content.append(page.extract_text())
            
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text content from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def extract_json_from_file(self, file_path: str) -> Dict[str, Any]:
        """Extract JSON content from file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except Exception as e:
            raise Exception(f"Error reading JSON file: {str(e)}")
    
    def extract_text_from_md(self, file_path: str) -> str:
        """Extract text content from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading Markdown file: {str(e)}")
    
    def extract_content(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Extract content based on file type."""
        content = {
            'text': '',
            'structured_data': None,
            'error': None
        }
        
        try:
            if file_extension == '.docx':
                content['text'] = self.extract_text_from_docx(file_path)
            elif file_extension == '.pdf':
                content['text'] = self.extract_text_from_pdf(file_path)
            elif file_extension == '.txt':
                content['text'] = self.extract_text_from_txt(file_path)
            elif file_extension == '.json':
                content['structured_data'] = self.extract_json_from_file(file_path)
                content['text'] = json.dumps(content['structured_data'], indent=2)
            elif file_extension == '.md':
                content['text'] = self.extract_text_from_md(file_path)
            
        except Exception as e:
            content['error'] = str(e)
        
        return content

    def cleanup_temp_files(self):
        """Clean up temporary files."""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception:
                pass  # Ignore cleanup errors
        self.temp_files.clear()

# Standalone main function (if needed, outside the class)
def main():
    st.title("BRD System")
    st.header("📄 Upload BRD Document")
    # Placeholder for upload_markdown_file, render_markdown_content, integrate_with_brd_validation
    st.info("Upload and processing logic goes here. Please implement upload_markdown_file, render_markdown_content, and integrate_with_brd_validation as needed.")